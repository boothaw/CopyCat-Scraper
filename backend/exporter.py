from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Map HTML heading tags to Word heading styles
HEADING_STYLE = {
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "h4": "Heading 4",
    "h5": "Heading 5",
    "h6": "Heading 6",
}


def _add_html_paragraph(doc: Document, element) -> None:
    """Convert a single BS4 block element into a Word paragraph."""
    tag = element.name if hasattr(element, "name") else None
    text = element.get_text(separator=" ", strip=True)
    if not text:
        return

    if tag in HEADING_STYLE:
        doc.add_heading(text, level=int(tag[1]))
    elif tag in ("ul", "ol"):
        for li in element.find_all("li", recursive=False):
            li_text = li.get_text(separator=" ", strip=True)
            if li_text:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(li_text)
    elif tag == "blockquote":
        p = doc.add_paragraph(style="Quote") if "Quote" in [s.name for s in doc.styles] else doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
    else:
        if text:
            doc.add_paragraph(text)


def build_docx(articles: list[dict]) -> BytesIO:
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i, article in enumerate(articles):
        if i > 0:
            doc.add_page_break()

        # Title
        title_para = doc.add_heading(article.get("title", "Untitled"), level=1)

        # Metadata line: URL and date
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
        url_run = meta.add_run(article.get("url", ""))
        url_run.italic = True
        url_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        url_run.font.size = Pt(9)

        date_val = article.get("date", "")
        if date_val:
            date_run = meta.add_run(f"  ·  {date_val}")
            date_run.italic = True
            date_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            date_run.font.size = Pt(9)

        doc.add_paragraph()  # spacer

        # Body content
        soup = BeautifulSoup(article.get("content", ""), "html.parser")

        block_tags = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "blockquote", "div"}

        def process_node(node):
            if not hasattr(node, "name") or node.name is None:
                # Bare text node — add as paragraph if meaningful
                text = str(node).strip()
                if text:
                    doc.add_paragraph(text)
                return

            if node.name in block_tags:
                if node.name == "div":
                    # Recurse into divs rather than treating them as blocks
                    for child in node.children:
                        process_node(child)
                else:
                    _add_html_paragraph(doc, node)
            # Inline elements at the top level — collect text
            else:
                text = node.get_text(separator=" ", strip=True)
                if text:
                    doc.add_paragraph(text)

        for child in soup.children:
            process_node(child)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
